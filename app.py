import streamlit as st
import assemblyai as aai
from pydub import AudioSegment
import os
from docx import Document
from io import BytesIO
from tqdm import tqdm

# -------------------------------
# Helper Functions
# -------------------------------
def split_audio(file_path, chunk_length_ms=300000):  # 5 minutes = 300,000 ms
    """Splits long audio into smaller chunks"""
    audio = AudioSegment.from_file(file_path)
    chunks = [audio[i:i + chunk_length_ms] for i in range(0, len(audio), chunk_length_ms)]
    chunk_files = []
    for idx, chunk in enumerate(chunks):
        chunk_name = f"chunk_{idx}.mp3"
        chunk.export(chunk_name, format="mp3")
        chunk_files.append(chunk_name)
    return chunk_files

def transcribe_audio(api_key, audio_file):
    """Handles transcription using AssemblyAI"""
    aai.settings.api_key = api_key
    config = aai.TranscriptionConfig(speech_model=aai.SpeechModel.universal)
    transcriber = aai.Transcriber(config=config)
    return transcriber.transcribe(audio_file)

def summarize_text(text, max_words=60):
    """Generates a short summary from full transcript"""
    words = text.split()
    if len(words) <= max_words:
        return text
    return " ".join(words[:max_words]) + "..."

def save_to_docx(text):
    """Saves full transcript as a .docx file"""
    doc = Document()
    doc.add_heading("Transcription", level=1)
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -------------------------------
# Streamlit UI
# -------------------------------
st.set_page_config(page_title="Smart Audio Transcriber", layout="centered")

st.title("🎧 Smart Audio Transcriber")
st.markdown("""
Upload your **MP3** or **WAV** voice note, and the app will transcribe it using AssemblyAI.  
It automatically handles large files by splitting them into smaller chunks.
""")

api_key = st.text_input("🔑 Enter your AssemblyAI API Key", type="password")
uploaded_file = st.file_uploader("📤 Upload your audio file", type=["mp3", "wav"])

if api_key and uploaded_file:
    try:
        # ✅ Validate file extension
        file_ext = os.path.splitext(uploaded_file.name)[1].lower()
        if file_ext not in [".mp3", ".wav"]:
            raise ValueError("Unsupported file type")

        # ✅ Save uploaded file temporarily
        file_path = uploaded_file.name
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        st.success(f"✅ File uploaded successfully: {uploaded_file.name}")

        # Check file size (in MB)
        file_size = os.path.getsize(file_path) / (1024 * 1024)
        st.info(f"File size: {file_size:.2f} MB")

        if st.button("🚀 Start Transcription"):
            st.write("Processing... Please wait ⏳")

            if file_size > 25:
                st.warning("Large file detected! Splitting into chunks for efficient processing.")
                chunk_files = split_audio(file_path)
                st.write(f"Split into {len(chunk_files)} chunks.")

                progress_bar = st.progress(0)
                all_text = ""
                for i, chunk_file in enumerate(chunk_files):
                    transcript = transcribe_audio(api_key, chunk_file)
                    if transcript.status == "error":
                        st.error(f"Error in chunk {i}: {transcript.error}")
                        break
                    all_text += transcript.text + "\n"
                    progress_bar.progress((i + 1) / len(chunk_files))
            else:
                transcript = transcribe_audio(api_key, file_path)
                if transcript.status == "error":
                    st.error(f"Transcription failed: {transcript.error}")
                else:
                    all_text = transcript.text

            # ✅ Display transcription results
            if "all_text" in locals():
                short_summary = summarize_text(all_text)
                st.subheader("📝 Short Summary")
                st.write(short_summary)

                st.subheader("📃 Full Transcript")
                with st.expander("Click to view full transcription"):
                    st.text_area("", all_text, height=250)

                docx_file = save_to_docx(all_text)
                st.download_button(
                    label="⬇️ Download Transcript (Word)",
                    data=docx_file,
                    file_name="transcript.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                st.success("✅ Transcription completed successfully!")

    except ValueError:
        st.error("❌ Unsupported audio format. Please upload only MP3 or WAV files.")
    except Exception as e:
        st.error(f"⚠️ An unexpected error occurred: {str(e)}")

